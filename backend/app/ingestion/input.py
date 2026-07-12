"""Lớp chuẩn hóa input (M4) — MỌI loại đầu vào quy về text rồi vào chung
pipeline analyze → classify → simulate → validate (M4 §1, §6).

Không loại input nào có pipeline riêng, không loại nào bypass validation.
Ảnh đi qua bước phiên dịch (Gemini Vision) thành text — Vision KHÔNG bao giờ
trả thẳng envelope (M4 §5). Code người dùng CHỈ được đọc, không execute (§8).
"""

from __future__ import annotations

import base64
import binascii
import io
import zipfile

from app.ai.gemini import call_gemini, load_skill

# ── Giới hạn ──────────────────────────────────────────────────
MAX_TEXT_CHARS = 8000
MAX_CODE_CHARS = 30_000
MAX_DOCX_BYTES = 2 * 1024 * 1024
MAX_IMAGE_BYTES = 4 * 1024 * 1024

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
VALID_IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp"}

# Magic bytes để phát hiện file giả đuôi/mime (§3, §5)
_IMAGE_SIGNATURES = {
    "image/png": [b"\x89PNG\r\n\x1a\n"],
    "image/jpeg": [b"\xff\xd8\xff"],
    "image/webp": [b"RIFF"],  # RIFF....WEBP
}

# Ngôn ngữ code nhận diện từ đuôi (§4) — trước mắt ưu tiên .py
_CODE_LANGS = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".c": "c",
    ".cpp": "cpp",
    ".java": "java",
    ".pas": "pascal",
}


class IngestError(Exception):
    """Lỗi chuẩn hóa input — main.py bắt và trả 400 với thông điệp tiếng Việt."""


def _decode_base64(content: str, limit: int, what: str) -> bytes:
    try:
        raw = base64.b64decode(content, validate=True)
    except (binascii.Error, ValueError):
        raise IngestError(f"{what} không phải dữ liệu base64 hợp lệ.")
    if len(raw) == 0:
        raise IngestError(f"{what} rỗng.")
    if len(raw) > limit:
        raise IngestError(f"{what} quá lớn (tối đa {limit // (1024 * 1024)}MB).")
    return raw


def _language_from_filename(filename: str | None) -> str:
    if not filename or "." not in filename:
        return "python"  # trước mắt mặc định .py
    ext = filename[filename.rindex(".") :].lower()
    return _CODE_LANGS.get(ext, "text")


def _extract_docx_text(raw: bytes) -> str:
    """Trích text từ .docx (§3): paragraph + text trong bảng, chuẩn hóa khoảng
    trắng, reject file hỏng. Dùng python-docx; không gửi binary cho LLM."""
    from docx import Document  # import trễ để test không cần thư viện nếu không chạm docx
    from docx.opc.exceptions import PackageNotFoundError

    # File giả đuôi: .docx thật là zip (bắt đầu bằng "PK")
    if not raw.startswith(b"PK"):
        raise IngestError("File .docx không hợp lệ (sai định dạng hoặc đã hỏng).")

    try:
        document = Document(io.BytesIO(raw))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError):
        raise IngestError("Không đọc được file .docx (có thể đã hỏng hoặc không đúng định dạng Word).")

    lines: list[str] = []
    for para in document.paragraphs:
        text = " ".join(para.text.split())
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    result = "\n".join(lines).strip()
    if not result:
        raise IngestError("File .docx không có nội dung văn bản nào đọc được.")
    return result[:MAX_TEXT_CHARS]


def _check_image(raw: bytes, mime_type: str | None) -> str:
    if mime_type not in VALID_IMAGE_MIMES:
        raise IngestError(
            f'Định dạng ảnh "{mime_type}" không được hỗ trợ. Chỉ nhận PNG, JPEG hoặc WEBP.'
        )
    signatures = _IMAGE_SIGNATURES[mime_type]
    if not any(raw.startswith(sig) for sig in signatures):
        raise IngestError("Nội dung ảnh không khớp định dạng khai báo (file có thể bị giả đuôi).")
    if mime_type == "image/webp" and not (len(raw) >= 12 and raw[8:12] == b"WEBP"):
        raise IngestError("Ảnh WEBP không hợp lệ.")
    return mime_type


async def ingest_to_text(
    input_type: str,
    content: str,
    filename: str | None,
    mime_type: str | None,
    api_key: str | None,
) -> str:
    """Chuẩn hóa InputPayload về text để đưa vào pipeline.

    - text: dùng thẳng.
    - code: đọc trực tiếp, giữ nguyên xuống dòng/indent, bọc trong khối
      ```<lang>``` để analyze hiểu đây là code. KHÔNG execute (§4, §8).
    - document (.docx): python-docx trích text (§3).
    - image: Gemini Vision phiên dịch thành text rồi mới vào pipeline (§5).

    Ném IngestError (→ 400) khi input sai; cần api_key riêng cho image.
    """
    if input_type == "text":
        text = content.strip()
        if len(text) > MAX_TEXT_CHARS:
            raise IngestError(f"Đề bài quá dài (tối đa {MAX_TEXT_CHARS} ký tự).")
        return text

    if input_type == "code":
        if len(content) > MAX_CODE_CHARS:
            raise IngestError(f"File code quá lớn (tối đa {MAX_CODE_CHARS} ký tự).")
        code = content.rstrip("\n")
        if not code.strip():
            raise IngestError("File code rỗng.")
        lang = _language_from_filename(filename)
        name = f" ({filename})" if filename else ""
        return f"Đoạn chương trình{name} cần phân tích:\n```{lang}\n{code}\n```"

    if input_type == "document":
        raw = _decode_base64(content, MAX_DOCX_BYTES, "File tài liệu")
        return _extract_docx_text(raw)

    if input_type == "image":
        raw = _decode_base64(content, MAX_IMAGE_BYTES, "Ảnh")
        valid_mime = _check_image(raw, mime_type)
        if not api_key:
            raise IngestError("__NEED_KEY__")  # main.py chuyển thành 503
        transcribed = await call_gemini(
            api_key,
            load_skill("transcribe"),
            "Chép lại nội dung đề bài trong ảnh này.",
            temperature=0.0,
            image={"mime_type": valid_mime, "data": content},
        )
        text = transcribed.strip()
        if len(text) < 10:
            raise IngestError("Không đọc được đề bài trong ảnh. Hãy chụp rõ hơn hoặc gõ tay đề.")
        return text[:MAX_TEXT_CHARS]

    raise IngestError(f'Loại đầu vào "{input_type}" không được hỗ trợ.')
