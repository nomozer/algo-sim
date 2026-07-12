# -*- coding: utf-8 -*-
"""Test lớp chuẩn hóa input (M4 §13).

Mọi loại input quy về text; không loại nào bypass pipeline; code KHÔNG execute.
Nhánh gọi Gemini Vision được mock — không cần key/mạng.
"""

import asyncio
import base64
import io

import pytest

from app.ingestion import input as ingest
from app.ingestion.input import IngestError, ingest_to_text

PNG_HEADER = b"\x89PNG\r\n\x1a\n"


def _run(coro):
    return asyncio.run(coro)


# ── text ──────────────────────────────────────────────────────

def test_text_giu_nguyen():
    out = _run(ingest_to_text("text", "  Tìm số lớn nhất trong dãy 3 1 2  ", None, None, None))
    assert out == "Tìm số lớn nhất trong dãy 3 1 2"


def test_text_qua_dai_bi_reject():
    with pytest.raises(IngestError, match="quá dài"):
        _run(ingest_to_text("text", "x" * 9000, None, None, None))


# ── code (.py) ────────────────────────────────────────────────

def test_code_giu_nguyen_indentation():
    code = "def f(a):\n    if a > 0:\n        return a\n    return 0\n"
    out = _run(ingest_to_text("code", code, "sol.py", None, None))
    # Indentation 4/8 space phải còn nguyên trong khối code
    assert "    if a > 0:" in out
    assert "        return a" in out
    assert "```python" in out
    assert "(sol.py)" in out


def test_code_nhan_dien_ngon_ngu_tu_duoi():
    out = _run(ingest_to_text("code", "let x = 1;", "a.ts", None, None))
    assert "```typescript" in out


def test_code_rong_bi_reject():
    with pytest.raises(IngestError, match="rỗng"):
        _run(ingest_to_text("code", "   \n  ", "a.py", None, None))


def test_code_khong_bi_execute(monkeypatch):
    """§8: code người dùng chỉ được đọc. ingest không được đụng eval/exec/subprocess."""
    import subprocess

    def _boom(*a, **k):
        raise AssertionError("subprocess bị gọi khi xử lý code người dùng!")

    monkeypatch.setattr(subprocess, "run", _boom)
    monkeypatch.setattr(subprocess, "Popen", _boom)
    # Code có payload nguy hiểm vẫn chỉ được bọc thành text, không chạy
    out = _run(ingest_to_text("code", "import os\nos.system('rm -rf /')\n", "x.py", None, None))
    assert "os.system" in out  # nằm nguyên trong text, không được thực thi


# ── document (.docx) ──────────────────────────────────────────

def _make_docx(paragraphs, table_rows=None) -> str:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, val in enumerate(row):
                t.rows[i].cells[j].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


def test_docx_trich_paragraph_va_bang():
    b64 = _make_docx(
        ["Tìm bạn điểm cao nhất tổ 1.", "Điểm số như sau:"],
        table_rows=[["An", "7.5"], ["Bình", "9"]],
    )
    out = _run(ingest_to_text("document", b64, "de.docx", None, None))
    assert "Tìm bạn điểm cao nhất tổ 1." in out
    assert "Điểm số như sau:" in out
    assert "An | 7.5" in out  # ô bảng nối bằng " | "
    assert "Bình | 9" in out


def test_docx_hong_bi_reject():
    fake = base64.b64encode(b"day khong phai file docx").decode()
    with pytest.raises(IngestError, match="không hợp lệ|không đọc được"):
        _run(ingest_to_text("document", fake, "gia.docx", None, None))


def test_docx_base64_sai_bi_reject():
    with pytest.raises(IngestError, match="base64"):
        _run(ingest_to_text("document", "@@@khong-phai-base64@@@", "a.docx", None, None))


def test_docx_rong_bi_reject():
    b64 = _make_docx([""])  # không paragraph nào có nội dung
    with pytest.raises(IngestError, match="không có nội dung"):
        _run(ingest_to_text("document", b64, "a.docx", None, None))


# ── image ─────────────────────────────────────────────────────

def test_image_mime_sai_bi_reject():
    b64 = base64.b64encode(PNG_HEADER + b"data").decode()
    with pytest.raises(IngestError, match="không được hỗ trợ"):
        _run(ingest_to_text("image", b64, "a.gif", "image/gif", None))


def test_image_gia_duoi_bi_reject():
    """Mime khai báo png nhưng nội dung không có magic bytes png."""
    b64 = base64.b64encode(b"khong-phai-png").decode()
    with pytest.raises(IngestError, match="không khớp định dạng"):
        _run(ingest_to_text("image", b64, "gia.png", "image/png", None))


def test_image_qua_lon_bi_reject():
    big = base64.b64encode(PNG_HEADER + b"x" * (5 * 1024 * 1024)).decode()
    with pytest.raises(IngestError, match="quá lớn"):
        _run(ingest_to_text("image", big, "big.png", "image/png", None))


def test_image_thieu_key_bao_hieu_need_key():
    b64 = base64.b64encode(PNG_HEADER + b"data").decode()
    with pytest.raises(IngestError, match="__NEED_KEY__"):
        _run(ingest_to_text("image", b64, "a.png", "image/png", None))


def test_image_phien_dich_thanh_text(monkeypatch):
    """§5: ảnh hợp lệ → Vision phiên dịch thành text; KHÔNG trả thẳng envelope."""

    async def fake_gemini(api_key, system_prompt, user_text, response_schema=None,
                          temperature=0.2, image=None):
        assert image is not None  # đúng là gọi vision với part ảnh
        assert image["mime_type"] == "image/png"
        return "Cho dãy 7, 9, 6. Tìm phần tử lớn nhất."

    monkeypatch.setattr(ingest, "call_gemini", fake_gemini)
    b64 = base64.b64encode(PNG_HEADER + b"data").decode()
    out = _run(ingest_to_text("image", b64, "de.png", "image/png", "khoa-gia"))
    assert out == "Cho dãy 7, 9, 6. Tìm phần tử lớn nhất."


def test_loai_input_la_bi_reject():
    with pytest.raises(IngestError):
        _run(ingest_to_text("video", "x", None, None, None))
